from __future__ import annotations

import csv
import re
from dataclasses import dataclass, field
from io import BytesIO, StringIO
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader


@dataclass
class ParsedBlock:
    text: str
    page_number: int | None = None
    title_path: str | None = None
    section_name: str | None = None
    block_type: str = "paragraph"
    metadata: dict = field(default_factory=dict)


def parse_document(file_bytes: bytes, file_name: str, file_ext: str) -> list[ParsedBlock]:
    ext = file_ext.lower().lstrip(".")
    if ext == "pdf":
        return parse_pdf(file_bytes, file_name)
    if ext == "docx":
        return parse_docx(file_bytes, file_name)
    if ext == "txt":
        return parse_txt(file_bytes, file_name)
    if ext == "md":
        return parse_markdown(file_bytes, file_name)
    if ext == "csv":
        return parse_csv(file_bytes, file_name)
    raise ValueError(f"Unsupported file extension: {file_ext}")


def clean_text(value: str) -> str:
    normalized = value.replace("\u3000", " ")
    normalized = re.sub(r"[ \t\r\f\v]+", " ", normalized)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()


def decode_text(file_bytes: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="replace")


def parse_pdf(file_bytes: bytes, file_name: str) -> list[ParsedBlock]:
    reader = PdfReader(BytesIO(file_bytes))
    blocks: list[ParsedBlock] = []
    for index, page in enumerate(reader.pages, start=1):
        text = clean_text(page.extract_text() or "")
        if text:
            blocks.append(
                ParsedBlock(
                    text=text,
                    page_number=index,
                    block_type="page",
                    metadata={"file_name": file_name},
                )
            )
    return blocks


def parse_docx(file_bytes: bytes, file_name: str) -> list[ParsedBlock]:
    document = DocxDocument(BytesIO(file_bytes))
    headings: list[str] = []
    blocks: list[ParsedBlock] = []

    for paragraph in document.paragraphs:
        text = clean_text(paragraph.text)
        if not text:
            continue

        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.startswith("Heading"):
            level = _heading_level(style_name)
            headings = headings[: level - 1]
            headings.append(text)
            blocks.append(
                ParsedBlock(
                    text=text,
                    title_path=" / ".join(headings),
                    section_name=text,
                    block_type="heading",
                    metadata={"file_name": file_name, "style": style_name},
                )
            )
        else:
            blocks.append(
                ParsedBlock(
                    text=text,
                    title_path=" / ".join(headings) or None,
                    section_name=headings[-1] if headings else None,
                    block_type="paragraph",
                    metadata={"file_name": file_name, "style": style_name},
                )
            )

    for table_index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [clean_text(cell.text) for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            blocks.append(
                ParsedBlock(
                    text="\n".join(rows),
                    title_path=" / ".join(headings) or None,
                    section_name=headings[-1] if headings else None,
                    block_type="table",
                    metadata={"file_name": file_name, "table_index": table_index},
                )
            )

    return blocks


def parse_txt(file_bytes: bytes, file_name: str) -> list[ParsedBlock]:
    text = decode_text(file_bytes)
    paragraphs = [clean_text(part) for part in re.split(r"\n\s*\n", text)]
    return [
        ParsedBlock(text=paragraph, block_type="paragraph", metadata={"file_name": file_name})
        for paragraph in paragraphs
        if paragraph
    ]


def parse_markdown(file_bytes: bytes, file_name: str) -> list[ParsedBlock]:
    text = decode_text(file_bytes)
    headings: list[str] = []
    buffer: list[str] = []
    blocks: list[ParsedBlock] = []

    def flush_buffer() -> None:
        content = clean_text("\n".join(buffer))
        if content:
            blocks.append(
                ParsedBlock(
                    text=content,
                    title_path=" / ".join(headings) or None,
                    section_name=headings[-1] if headings else None,
                    block_type="paragraph",
                    metadata={"file_name": file_name},
                )
            )
        buffer.clear()

    for line in text.splitlines():
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading_match:
            flush_buffer()
            level = len(heading_match.group(1))
            heading = clean_text(heading_match.group(2))
            headings = headings[: level - 1]
            headings.append(heading)
            blocks.append(
                ParsedBlock(
                    text=heading,
                    title_path=" / ".join(headings),
                    section_name=heading,
                    block_type="heading",
                    metadata={"file_name": file_name, "level": level},
                )
            )
        else:
            buffer.append(line)

    flush_buffer()
    return blocks


def parse_csv(file_bytes: bytes, file_name: str) -> list[ParsedBlock]:
    text = decode_text(file_bytes)
    sample = text[:2048]
    try:
        dialect = csv.Sniffer().sniff(sample)
    except csv.Error:
        dialect = csv.excel

    reader = csv.reader(StringIO(text), dialect)
    rows = list(reader)
    if not rows:
        return []

    header = rows[0]
    data_rows = rows[1:] if header else rows
    blocks: list[ParsedBlock] = []
    batch_size = 20

    for start in range(0, len(data_rows), batch_size):
        batch = data_rows[start : start + batch_size]
        lines = [f"Columns: {', '.join(header)}"] if header else []
        for row_index, row in enumerate(batch, start=start + 1):
            cells = []
            for column_index, value in enumerate(row):
                column_name = header[column_index] if column_index < len(header) else f"column_{column_index + 1}"
                cells.append(f"{column_name}: {clean_text(value)}")
            lines.append(f"Row {row_index}: " + "; ".join(cells))

        blocks.append(
            ParsedBlock(
                text="\n".join(lines),
                block_type="table",
                metadata={
                    "file_name": file_name,
                    "row_start": start + 1,
                    "row_end": start + len(batch),
                    "source_ext": Path(file_name).suffix.lower(),
                },
            )
        )

    return blocks


def _heading_level(style_name: str) -> int:
    match = re.search(r"(\d+)$", style_name)
    if not match:
        return 1
    return max(1, min(6, int(match.group(1))))
